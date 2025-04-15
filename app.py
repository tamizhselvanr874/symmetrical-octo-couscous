import os
from openai import AzureOpenAI
import json
import re

def get_azure_client():
    """Initialize and return the Azure OpenAI client."""
    azure_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
    api_key = os.getenv("AZURE_OPENAI_API_KEY")
    
    client = AzureOpenAI(
        azure_endpoint=azure_endpoint,
        api_key=api_key,
        api_version="2024-10-01-preview",
    )
    return client

def validate_trademark_relevance(conflicts_array, proposed_goods_services):
    """
    Pre-filter trademarks that don't have similar or identical goods/services
    This function is implemented in code rather than relying on GPT
    
    Args:
        conflicts_array: List of trademark conflicts
        proposed_goods_services: Goods/services of the proposed trademark
        
    Returns:
        filtered_conflicts: List of relevant trademark conflicts
        excluded_count: Number of trademarks excluded
    """
    # Parse conflicts_array if it's a string (assuming JSON format)
    if isinstance(conflicts_array, str):
        try:
            conflicts = json.loads(conflicts_array)
        except json.JSONDecodeError:
            # If it's not valid JSON, try to parse it as a list of dictionaries
            conflicts = eval(conflicts_array) if conflicts_array.strip().startswith("[") else []
    else:
        conflicts = conflicts_array
    
    # Initialize lists for relevant and excluded trademarks
    relevant_conflicts = []
    excluded_count = 0
    
    # Define a function to check similarity between goods/services
    def is_similar_goods_services(existing_goods, proposed_goods):
        # Convert to lowercase for case-insensitive comparison
        existing_lower = existing_goods.lower()
        proposed_lower = proposed_goods.lower()
        
        # Check for exact match
        if existing_lower == proposed_lower:
            return True
        
        # Check if one contains the other
        if existing_lower in proposed_lower or proposed_lower in existing_lower:
            return True
        
        # Check for overlapping keywords
        # Extract significant keywords from both descriptions
        existing_keywords = set(re.findall(r'\b\w+\b', existing_lower))
        proposed_keywords = set(re.findall(r'\b\w+\b', proposed_lower))
        
        # Remove common stop words
        stop_words = {'and', 'or', 'the', 'a', 'an', 'in', 'on', 'for', 'of', 'to', 'with'}
        existing_keywords = existing_keywords - stop_words
        proposed_keywords = proposed_keywords - stop_words
        
        # Calculate keyword overlap
        if len(existing_keywords) > 0 and len(proposed_keywords) > 0:
            overlap = len(existing_keywords.intersection(proposed_keywords))
            overlap_ratio = overlap / min(len(existing_keywords), len(proposed_keywords))
            
            # If significant overlap (more than 30%), consider them similar
            if overlap_ratio > 0.3:
                return True
        
        return False
    
    # Process each conflict
    for conflict in conflicts:
        # Ensure conflict has goods/services field
        if 'goods_services' in conflict:
            if is_similar_goods_services(conflict['goods_services'], proposed_goods_services):
                relevant_conflicts.append(conflict)
            else:
                excluded_count += 1
        else:
            # If no goods/services field, include it for safety
            relevant_conflicts.append(conflict)
    
    return relevant_conflicts, excluded_count

def filter_by_gpt_response(conflicts, gpt_json):
    """
    Removes trademarks that GPT flagged as lacking goods/services overlap.
    
    Args:
        conflicts: Original list of trademark conflicts
        gpt_json: JSON object from GPT with 'results' key
    
    Returns:
        Filtered list of conflicts that GPT identified as overlapping
    """
    # Parse the GPT response if it's a string
    if isinstance(gpt_json, str):
        try:
            gpt_json = json.loads(gpt_json)
        except json.JSONDecodeError:
            # If JSON is invalid, keep original conflicts
            return conflicts
    
    gpt_results = gpt_json.get("results", [])
    
    # Build a set of marks with overlap for quick membership checking
    overlapping_marks = {
        result["mark"]
        for result in gpt_results
        if result.get("overlap") is True
    }
    
    # Retain conflicts only if they appear in overlapping_marks
    filtered_conflicts = [
        c for c in conflicts
        if c.get("mark") in overlapping_marks
    ]
    
    return filtered_conflicts

def clean_and_format_opinion(comprehensive_opinion, json_data=None):
    """
    Process the comprehensive trademark opinion to:
    1. Maintain comprehensive listing of all relevant trademark hits
    2. Remove duplicated content while preserving all unique trademark references
    3. Format the opinion for better readability
    4. Ensure consistent structure with clear sections
    
    Args:
        comprehensive_opinion: Raw comprehensive opinion from previous steps
        json_data: Optional structured JSON data from previous steps
        
    Returns:
        A cleaned, formatted, and optimized trademark opinion
    """
    client = get_azure_client()
    
    system_prompt = """
    You are a trademark attorney specializing in clear, comprehensive trademark opinions.
    
    FORMAT THE TRADEMARK OPINION USING THE EXACT STRUCTURE PROVIDED BELOW:
    
    ```
REFINED TRADEMARK OPINION: [MARK NAME]
Class: [Class Number]
Goods and Services: [Goods/Services Description]

Section I: Comprehensive Trademark Hit Analysis
(a) Identical Marks:
| Trademark | Owner | Goods & Services | Status | Class | Class Match | Goods & Services Match |
|------------|--------|------------------|--------|------|-------------|------------------------|
| [Mark 1] | [Owner] | [Goods/Services] | [Status] | [Class] | [True/False] | [True/False] |

(b) One Letter and Two Letter Analysis:
| Trademark | Owner | Goods & Services | Status | Class | Difference Type | Class Match | Goods & Services Match |
|------------|--------|------------------|--------|------|----------------|-------------|------------------------|
| [Mark 1] | [Owner] | [Goods/Services] | [Status] | [Class] | [One/Two Letter] | [True/False] | [True/False] |

(c) Phonetically, Semantically & Functionally Similar Analysis:
| Trademark | Owner | Goods & Services | Status | Class | Similarity Type | Class Match | Goods & Services Match |
|------------|--------|------------------|--------|------|-----------------|-------------|------------------------|
| [Mark 1] | [Owner] | [Goods/Services] | [Status] | [Class] | [Phonetic/Semantic/Functional] | [True/False] | [True/False] |

Section II: Component Analysis
(a) Component Analysis:

Component 1: [First Component]
| Trademark | Owner | Goods & Services | Status | Class | Class Match | Goods & Services Match |
|-----------|--------|------------------|--------|-------|-------------|------------------------|
| [Mark 1] | [Owner] | [Goods/Services] | [Status] | [Class] | [True/False] | [True/False] |

Component A: [Second Component]
| Trademark | Owner | Goods & Services | Status | Class | Class Match | Goods & Services Match |
|-----------|--------|------------------|--------|-------|-------------|------------------------|
| [Mark 1] | [Owner] | [Goods/Services] | [Status] | [Class] | [True/False] | [True/False] |

(b) Crowded Field Analysis:
- **Total compound mark hits found**: [NUMBER]
- **Marks with different owners**: [NUMBER] ([PERCENTAGE]%)
- **Crowded Field Status**: [YES/NO]
- **Analysis**: 
  [DETAILED EXPLANATION OF FINDINGS INCLUDING RISK IMPLICATIONS IF FIELD IS CROWDED]

Section III: Risk Assessment and Summary

Descriptiveness:
- [KEY POINT ABOUT DESCRIPTIVENESS]

Aggressive Enforcement and Litigious Behavior:
- **Known Aggressive Owners**:
  * [Owner 1]: [Enforcement patterns]
  * [Owner 2]: [Enforcement patterns]
- **Enforcement Landscape**:
  * [KEY POINT ABOUT ENFORCEMENT LANDSCAPE]
  * [ADDITIONAL POINT ABOUT ENFORCEMENT LANDSCAPE]

Risk Category for Registration:
- **[REGISTRATION RISK LEVEL: HIGH/MEDIUM-HIGH/MEDIUM/MEDIUM-LOW/LOW]**
- [EXPLANATION OF REGISTRATION RISK LEVEL WITH FOCUS ON CROWDED FIELD ANALYSIS]

Risk Category for Use:
- **[USE RISK LEVEL: HIGH/MEDIUM-HIGH/MEDIUM/MEDIUM-LOW/LOW]**
- [EXPLANATION OF USE RISK LEVEL]
    ```

    **IMPORTANT INSTRUCTIONS:**
    1. Maintain ALL unique trademark references from the original opinion.
    2. Present trademarks in clear, easy-to-read tables following the format above.
    3. Ensure ALL findings from the original opinion are preserved but avoid redundancy.
    4. Include owner names and goods/services details for each mark.
    5. Include trademark search exclusions in the summary section.
    6. Ensure the final opinion is comprehensive yet concise.
    7. For each section, include all relevant trademarks without omission.
    8. Maintain the exact structure provided above with clear section headings.
    9. For each mark, determine and include:
       - "Class Match" (True/False): Whether the mark's class exactly matches the proposed trademark's class OR is in a coordinated/related class group.
       - "Goods & Services Match" (True/False): Whether the mark's goods/services are similar to the proposed trademark's goods/services.
    10. Follow the specified structure exactly:
        - Section I focuses on overall hits, including One/Two Letter Analysis
        - Section II focuses only on component hits
        - In Section II, perform Crowded Field Analysis focusing on owner diversity
    11. State "None" when no results are found for a particular subsection
    12. Do NOT include recommendations in the summary
    13. Include aggressive enforcement analysis in Section III with details on any owners known for litigious behavior
    14. IMPORTANT: When assessing "Class Match", consider not only exact class matches but also coordinated or related classes based on the goods/services.
    15. NEVER replace full goods/services descriptions with just class numbers in the output tables. Always include the complete goods/services text.
    """
    
    # Send the original opinion to be reformatted
    user_message = f"""
    Please reformat the following comprehensive trademark opinion according to the refined structure:
    
    Proposed Trademark: {json_data.get('proposed_name', 'N/A')}
    Class: {json_data.get('proposed_class', 'N/A')}
    Goods and Services: {json_data.get('proposed_goods_services', 'N/A')}
    
    Original Opinion:
    {comprehensive_opinion}
    
    Follow the exact structure provided in the instructions, ensuring all trademark references are maintained.
    
    For each mark in the tables, you must evaluate and include:
    1. Owner name
    2. Goods & Services description - ALWAYS include the FULL goods/services text, not just class numbers
    3. Class Match (True/False): 
       - Mark True if the mark's class exactly matches the proposed class "{json_data.get('proposed_class', 'N/A')}"
       - ALSO mark True if the mark's class is in a coordinated or related class grouping with the proposed class
       - First identify all coordinated classes based on the proposed goods/services: "{json_data.get('proposed_goods_services', 'N/A')}"
       - Then mark True for any mark in those coordinated classes
    4. Goods & Services Match (True/False): Compare the mark's goods/services to the proposed goods/services "{json_data.get('proposed_goods_services', 'N/A')}" and mark True if they are semantically similar.
    
    IMPORTANT REMINDERS FOR CROWDED FIELD ANALYSIS:
    - Include exact counts and percentages for:
      * Total compound mark hits found
      * Number and percentage of marks with different owners
      * Crowded Field Status (YES if >50% have different owners)
    - Clearly explain risk implications if field is crowded
    - Section I should include ALL hits (overall hits), not just compound mark hits
    - Section II should focus ONLY on compound mark hits
    - One and Two Letter Analysis should ONLY be in Section I, not Section II
    - If no results are found for a particular subsection, state "None"
    - Do NOT include recommendations in the summary
    - Include aggressive enforcement analysis in Section III with details on any owners known for litigious behavior
    """
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message}
            ],
            temperature=0.0,
        )
        
        # Extract and return the formatted opinion
        if response.choices and len(response.choices) > 0:
            formatted_opinion = response.choices[0].message.content
            
            # Filter out rows where both "Class Match" and "Goods & Services Match" are False
            filtered_opinion = []
            for line in formatted_opinion.splitlines():
                if "|" in line:  # Check if the line is part of a table
                    parts = line.split("|")
                    if len(parts) >= 7:  # Ensure the line has enough columns
                        # Check if this is a header row by looking for specific column header text
                        if "Class Match" in line or "Trademark" in line:
                            filtered_opinion.append(line)
                        else:
                            # For data rows, check the Class Match and Goods & Services Match values
                            class_match_idx = -3  # Second to last column
                            goods_services_match_idx = -1  # Last column
                            
                            class_match = "true" in parts[class_match_idx].strip().lower()
                            goods_services_match = "true" in parts[goods_services_match_idx].strip().lower()
                            
                            if class_match or goods_services_match:
                                filtered_opinion.append(line)
                    else:
                        # Include table formatting lines and other table parts
                        filtered_opinion.append(line)
                else:
                    # Include all non-table lines
                    filtered_opinion.append(line)

            # Join the filtered lines back into a single string
            filtered_opinion = "\n".join(filtered_opinion)
            
            return filtered_opinion
        else:
            return "Error: No response received from the language model."
    except Exception as e:
        return f"Error during opinion formatting: {str(e)}"
        
  
def consistency_check(mark, results):
    """
    Consistency checking function to ensure accuracy of analysis results.
    
    Args:
        mark: The proposed trademark name
        results: Raw analysis results
        
    Returns:
        Corrected and validated results
    """
    corrected_results = results.copy()
    
    # Ensure all entries have required fields
    required_fields = ['mark', 'owner', 'goods_services', 'status', 'class', 'class_match', 'goods_services_match']
    
    # Check identical marks
    for i, item in enumerate(corrected_results.get('identical_marks', [])):
        # Validate that mark name is indeed identical
        if item.get('mark', '').lower() != mark.lower():
            # Remove from identical marks if not actually identical
            corrected_results['identical_marks'][i] = None
        
        # Ensure all required fields exist
        for field in required_fields:
            if field not in item:
                if field == 'class_match' or field == 'goods_services_match':
                    corrected_results['identical_marks'][i][field] = False
                else:
                    corrected_results['identical_marks'][i][field] = "Unknown"
    
    # Remove None entries
    corrected_results['identical_marks'] = [item for item in corrected_results.get('identical_marks', []) if item is not None]
    
    # Similar checks for one_letter_marks
    for i, item in enumerate(corrected_results.get('one_letter_marks', [])):
        # Validate actual one letter difference
        if not is_one_letter_difference(item.get('mark', ''), mark):
            corrected_results['one_letter_marks'][i] = None
            
        # Ensure all required fields exist
        for field in required_fields:
            if field not in item and field != 'difference_type':
                if field == 'class_match' or field == 'goods_services_match':
                    corrected_results['one_letter_marks'][i][field] = False
                else:
                    corrected_results['one_letter_marks'][i][field] = "Unknown"
    
    # Remove None entries
    corrected_results['one_letter_marks'] = [item for item in corrected_results.get('one_letter_marks', []) if item is not None]
    
    # Similar checks for two_letter_marks
    for i, item in enumerate(corrected_results.get('two_letter_marks', [])):
        # Validate actual two letter difference
        if not is_two_letter_difference(item.get('mark', ''), mark):
            corrected_results['two_letter_marks'][i] = None
            
        # Ensure all required fields exist
        for field in required_fields:
            if field not in item and field != 'difference_type':
                if field == 'class_match' or field == 'goods_services_match':
                    corrected_results['two_letter_marks'][i][field] = False
                else:
                    corrected_results['two_letter_marks'][i][field] = "Unknown"
    
    # Remove None entries
    corrected_results['two_letter_marks'] = [item for item in corrected_results.get('two_letter_marks', []) if item is not None]
    
    # Check similar_marks
    for i, item in enumerate(corrected_results.get('similar_marks', [])):
        # Ensure all required fields exist
        for field in required_fields:
            if field not in item and field != 'similarity_type':
                if field == 'class_match' or field == 'goods_services_match':
                    corrected_results['similar_marks'][i][field] = False
                else:
                    corrected_results['similar_marks'][i][field] = "Unknown"
    
    return corrected_results


def is_one_letter_difference(mark1, mark2):
    """
    Check if two marks have a one letter difference.
    
    Args:
        mark1: First mark
        mark2: Second mark
        
    Returns:
        Boolean indicating if there's a one letter difference
    """
    # Handle case-insensitivity
    mark1 = mark1.lower()
    mark2 = mark2.lower()
    
    # If length difference is greater than 1, definitely not a one letter difference
    if abs(len(mark1) - len(mark2)) > 1:
        return False
    
    # Count differences
    differences = 0
    
    # Same length - check for substitution
    if len(mark1) == len(mark2):
        for c1, c2 in zip(mark1, mark2):
            if c1 != c2:
                differences += 1
                if differences > 1:
                    return False
    # Different length - check for insertion/deletion
    else:
        # Make sure mark1 is the shorter one for simplicity
        if len(mark1) > len(mark2):
            mark1, mark2 = mark2, mark1
            
        i, j = 0, 0
        while i < len(mark1) and j < len(mark2):
            if mark1[i] != mark2[j]:
                # Skip this character in the longer string
                j += 1
                differences += 1
                if differences > 1:
                    return False
            else:
                i += 1
                j += 1
                
    return differences == 1


def is_two_letter_difference(mark1, mark2):
    """
    Check if two marks have a two letter difference.
    
    Args:
        mark1: First mark
        mark2: Second mark
        
    Returns:
        Boolean indicating if there's a two letter difference
    """
    # Handle case-insensitivity
    mark1 = mark1.lower()
    mark2 = mark2.lower()
    
    # If length difference is greater than 2, definitely not a two letter difference
    if abs(len(mark1) - len(mark2)) > 2:
        return False
    
    # Use Levenshtein distance for accurate measurement
    return levenshtein_distance(mark1, mark2) == 2


def levenshtein_distance(s1, s2):
    """
    Calculate the Levenshtein distance between two strings.
    This measures the minimum number of single-character edits needed to change one string into another.
    
    Args:
        s1: First string
        s2: Second string
        
    Returns:
        The edit distance between the strings
    """
    if len(s1) < len(s2):
        return levenshtein_distance(s2, s1)

    if len(s2) == 0:
        return len(s1)

    previous_row = range(len(s2) + 1)
    for i, c1 in enumerate(s1):
        current_row = [i + 1]
        for j, c2 in enumerate(s2):
            insertions = previous_row[j + 1] + 1
            deletions = current_row[j] + 1
            substitutions = previous_row[j] + (c1 != c2)
            current_row.append(min(insertions, deletions, substitutions))
        previous_row = current_row
    
    return previous_row[-1]


def section_one_analysis(mark, class_number, goods_services, relevant_conflicts):  
    """
    Perform Section I: Comprehensive Trademark Hit Analysis using chain of thought prompting.
    This approach explicitly walks through the analysis process to ensure consistent results.
    Includes phonetic and semantic similarity checks.
    """  
    client = get_azure_client()  
    
    # Helper function for semantic equivalence
    def is_semantically_equivalent(name1, name2, threshold=0.50):
        embeddings1 = semantic_model.encode(name1, convert_to_tensor=True)
        embeddings2 = semantic_model.encode(name2, convert_to_tensor=True)
        similarity_score = util.cos_sim(embeddings1, embeddings2).item()
        return similarity_score >= threshold

    # Helper function for phonetic equivalence
    def is_phonetically_equivalent(name1, name2, threshold=50):
        return fuzz.ratio(name1.lower(), name2.lower()) >= threshold

    system_prompt = """
    You are a trademark expert attorney specializing in trademark opinion writing. I need you to analyze potential trademark conflicts using chain of thought reasoning.
    
    First, I want you to think step by step:
    
    1. STEP 1 - COORDINATED CLASS ANALYSIS:
       a) Carefully examine the proposed goods/services: "{goods_services}"
       b) Based on these goods/services, identify which other trademark classes would be considered related or coordinated with the primary class {class_number}
       c) Document your reasoning for each coordinated class you identify
       d) Create a final list of all classes that should be considered for conflict analysis
    
    2. STEP 2 - IDENTICAL MARK ANALYSIS:
       a) Identify any marks that EXACTLY match the proposed trademark "{mark}" (case-insensitive)
       b) For each identical mark, verify:
          - Is it in the SAME class as the proposed mark?
          - Is it in a COORDINATED class you identified in Step 1?
          - Are the goods/services similar or related to the proposed goods/services?
       c) For each mark, explicitly determine class_match and goods_services_match values
    
    3. STEP 3 - ONE LETTER DIFFERENCE ANALYSIS:
       a) Identify marks that have EXACTLY ONE letter different from the proposed mark
       b) This can be through substitution (one letter different), addition (one extra letter), or deletion (one letter missing)
       c) For each mark, explicitly determine class_match and goods_services_match values
    
    4. STEP 4 - TWO LETTER DIFFERENCE ANALYSIS:
       a) Identify marks that have EXACTLY TWO letters different from the proposed mark
       b) This can be through substitution, addition, deletion, or a combination
       c) For each mark, explicitly determine class_match and goods_services_match values
    
    5. STEP 5 - SIMILAR MARK ANALYSIS:
       a) Identify marks that are similar in sound (phonetically), meaning (semantically), or function
       b) For phonetic similarity, consider marks that sound similar when spoken aloud (use the phonetic algorithm)
       c) For semantic similarity, consider marks that have similar meanings or connotations
       d) Explicitly state the similarity type (Phonetic/Semantic/Functional) for each similar mark
       e) For each mark, explicitly determine class_match and goods_services_match values
    
    6. STEP 6 - CROWDED FIELD ANALYSIS:
       a) Calculate the total number of potentially conflicting marks identified
       b) Calculate what percentage of these marks have different owners
       c) Determine if the field is "crowded" (>50% different owners)
       d) Explain the implications for trademark protection
    
    For each conflict you identify, include comprehensive details:
    - The exact mark name
    - The owner name
    - The full goods/services description (not just class numbers)
    - Registration status
    - Class number
    - Whether there's a class match (true/false)
    - Whether there's a goods/services match (true/false)
    - For similar marks, include similarity type (Phonetic/Semantic/Functional)
    
    YOUR RESPONSE MUST BE IN JSON FORMAT:
    {
      "identified_coordinated_classes": [LIST OF RELATED CLASS NUMBERS],
      "coordinated_classes_explanation": "[EXPLANATION OF WHY THESE CLASSES ARE RELATED TO THE PROPOSED TRADEMARK]",
      "identical_marks": [
        {
          "mark": "[TRADEMARK NAME]",
          "owner": "[OWNER NAME]",
          "goods_services": "[GOODS/SERVICES]",
          "status": "[LIVE/DEAD]",
          "class": "[CLASS]",
          "class_match": true|false,
          "goods_services_match": true|false
        }
      ],
      "one_letter_marks": [
        {
          "mark": "[TRADEMARK NAME]",
          "owner": "[OWNER NAME]",
          "goods_services": "[GOODS/SERVICES]",
          "status": "[LIVE/DEAD]",
          "class": "[CLASS]",
          "difference_type": "One Letter",
          "class_match": true|false,
          "goods_services_match": true|false
        }
      ],
      "two_letter_marks": [
        {
          "mark": "[TRADEMARK NAME]",
          "owner": "[OWNER NAME]",
          "goods_services": "[GOODS/SERVICES]",
          "status": "[LIVE/DEAD]",
          "class": "[CLASS]",
          "difference_type": "Two Letter",
          "class_match": true|false,
          "goods_services_match": true|false
        }
      ],
      "similar_marks": [
        {
          "mark": "[TRADEMARK NAME]",
          "owner": "[OWNER NAME]",
          "goods_services": "[GOODS/SERVICES]",
          "status": "[LIVE/DEAD]",
          "class": "[CLASS]",
          "similarity_type": "[Phonetic|Semantic|Functional]",
          "class_match": true|false,
          "goods_services_match": true|false
        }
      ],
      "crowded_field": {
        "is_crowded": true|false,
        "percentage": [PERCENTAGE],
        "explanation": "[EXPLANATION]"
      }
    }
""" 
  
    user_message = f""" 
    Proposed Trademark: {mark}
    Class: {class_number}
    Goods/Services: {goods_services}
    
    Trademark Conflicts:
    {json.dumps(relevant_conflicts, indent=2)}
    
    Analyze ONLY Section I: Comprehensive Trademark Hit Analysis. Walk through each step methodically:
    
    STEP 1: First, carefully analyze the proposed goods/services and identify ALL coordinated classes.
    STEP 2: Then identify EXACT matches to the trademark "{mark}"
    STEP 3: Next, identify marks with ONE letter difference (substitution, addition, or deletion)
    STEP 4: Then identify marks with TWO letter differences
    STEP 5: Finally, identify phonetically, semantically, or functionally similar marks
        - For phonetic similarity, carefully compare how the marks sound when spoken
        - For semantic similarity, analyze the meaning and connotations
        - Explicitly state the similarity type (Phonetic/Semantic/Functional) for each
    STEP 6: Perform crowded field analysis with precise calculations
    
    IMPORTANT REMINDERS:
    - Focus on matches to the ENTIRE trademark name, not just components
    - Include owner names and goods/services details for each mark
    - For Class Match (True/False):
      * First, explicitly identify all coordinated classes related to the proposed goods/services
      * Mark True if the mark's class exactly matches the proposed class "{class_number}"
      * ALSO mark True if the mark's class is in a coordinated or related class grouping you identified
    - For Goods & Services Match (True/False), compare the mark's goods/services to the proposed goods/services
    - Always include the FULL goods/services description in your output, not just the class number
    - For One/Two Letter differences, carefully verify the exact letter count difference
    - For Similar marks, you MUST explicitly state whether similarity is Phonetic, Semantic, or Functional
    - DO NOT omit any marks that are phonetically similar - they must appear in the similar_marks section with similarity_type="Phonetic"
"""  
  
    try:  
        response = client.chat.completions.create(  
            model="gpt-4o",  
            messages=[  
                {"role": "system", "content": system_prompt},  
                {"role": "user", "content": user_message}  
            ],  
            temperature=0.0,  
        )  
  
        if response.choices and len(response.choices) > 0:  
            content = response.choices[0].message.content  
  
            # Extract JSON data  
            json_match = re.search(r'```json\s*(.*?)\s*```|({[\s\S]*})', content, re.DOTALL)  
            if json_match:  
                json_str = json_match.group(1) or json_match.group(2)  
                try:  
                    raw_results = json.loads(json_str)  
                    # Apply consistency checking  
                    corrected_results = consistency_check(mark, raw_results)  
                    
                    # Additional validation for phonetic/semantic matches
                    similar_marks = corrected_results.get('similar_marks', [])
                    new_similar_marks = []
                    
                    # First process existing similar marks
                    for similar_mark in similar_marks:
                        conflict_mark = similar_mark['mark']
                        if similar_mark.get('similarity_type') == 'Phonetic':
                            similar_mark['valid_phonetic_match'] = is_phonetically_equivalent(mark, conflict_mark)
                        elif similar_mark.get('similarity_type') == 'Semantic':
                            similar_mark['valid_semantic_match'] = is_semantically_equivalent(mark, conflict_mark)
                        new_similar_marks.append(similar_mark)
                    
                    # Now check all conflicts for potential phonetic matches that might have been missed
                    for conflict in relevant_conflicts:
                        conflict_mark = conflict.get('trademark_name', '')
                        if is_phonetically_equivalent(mark, conflict_mark):
                            # Check if this conflict is already in similar_marks
                            already_listed = any(
                                sm['mark'] == conflict_mark 
                                and sm.get('similarity_type') == 'Phonetic' 
                                for sm in new_similar_marks
                            )
                            
                            if not already_listed:
                                # Add as a new phonetic match
                                new_similar_marks.append({
                                    'mark': conflict_mark,
                                    'owner': conflict.get('owner', 'Unknown'),
                                    'goods_services': conflict.get('goods_services', ''),
                                    'status': conflict.get('status', 'Unknown'),
                                    'class': conflict.get('class', ''),
                                    'similarity_type': 'Phonetic',
                                    'class_match': conflict.get('class', '') == class_number or conflict.get('class', '') in raw_results.get('identified_coordinated_classes', []),
                                    'goods_services_match': True,  # Assuming validate_trademark_relevance already filtered these
                                    'valid_phonetic_match': True,
                                    'added_by_validation': True  # Flag to indicate this was added in validation
                                })
                    
                    corrected_results['similar_marks'] = new_similar_marks
                    
                    return corrected_results  
                except json.JSONDecodeError:  
                    return {  
                        "identified_coordinated_classes": [],
                        "coordinated_classes_explanation": "Unable to identify coordinated classes",
                        "identical_marks": [],  
                        "one_letter_marks": [],  
                        "two_letter_marks": [],  
                        "similar_marks": [],
                        "crowded_field": {
                            "is_crowded": False,
                            "percentage": 0,
                            "explanation": "Unable to determine crowded field status"
                        }
                    }  
            else:  
                return {  
                    "identified_coordinated_classes": [],
                    "coordinated_classes_explanation": "Unable to identify coordinated classes",
                    "identical_marks": [],  
                    "one_letter_marks": [],  
                    "two_letter_marks": [],  
                    "similar_marks": [],
                    "crowded_field": {
                        "is_crowded": False,
                        "percentage": 0,
                        "explanation": "Unable to determine crowded field status"
                    }
                }  
        else:  
            return {  
                "identified_coordinated_classes": [],
                "coordinated_classes_explanation": "Unable to identify coordinated classes",
                "identical_marks": [],  
                "one_letter_marks": [],  
                "two_letter_marks": [],  
                "similar_marks": [],
                "crowded_field": {
                    "is_crowded": False,
                    "percentage": 0,
                    "explanation": "Unable to determine crowded field status"
                }
            }  
    except Exception as e:  
        print(f"Error in section_one_analysis: {str(e)}")  
        return {  
            "identified_coordinated_classes": [],
            "coordinated_classes_explanation": "Error occurred during analysis",
            "identical_marks": [],  
            "one_letter_marks": [],  
            "two_letter_marks": [],  
            "similar_marks": [],
            "crowded_field": {
                "is_crowded": False,
                "percentage": 0,
                "explanation": "Error occurred during analysis"
            }
        }
    

def section_two_analysis(mark, class_number, goods_services, relevant_conflicts):  
    """Perform Section II: Component Analysis."""  
    client = get_azure_client()  
  
    system_prompt = """
    You are a trademark expert attorney specializing in trademark opinion writing.
    
    Please perform an analysis focusing on Section II: Component Analysis. In this section, you should:

    (a) Identify and break the proposed trademark into its components (if it is compound).
    (b) For each component, analyze marks that incorporate that component.
    (c) For each conflict record, include details such as the owner, goods/services, registration status, and class information.
    (d) Determine flags for both "goods_services_match" and "class_match."

    IMPORTANT INSTRUCTIONS FOR COORDINATED CLASS ANALYSIS:
    • First, analyze the provided goods/services description to identify which trademark classes are closely related or coordinated with the proposed trademark's class.
    • You MUST thoroughly analyze and include conflicts across RELATED and COORDINATED classes, not just exact class matches.
    • Common coordinated class groupings include:
      - Food and beverage products: Consider classes 29, 30, 31, 32, 35, 43
      - Furniture and home goods: Consider classes 20, 35, 42
      - Clothing and fashion: Consider classes 18, 25, 35
      - Technology and software: Consider classes 9, 38, 42
      - Health and beauty: Consider classes 3, 5, 44
      - Entertainment: Consider classes 9, 41, 42
    • However, do not limit yourself to these examples - use your expertise to identify all relevant coordinated classes for the specific goods/services.
    • If ANY component of the proposed trademark appears in ANY other class, this must be flagged.
    • DO NOT MISS conflicts across coordinated classes - this is CRITICAL.

    • When comparing classes, do not only check for an exact match. Explicitly check whether a conflict is registered in a coordinated or related class; for instance, if a conflict is in a related class grouping, you must mark it with class_match = True.
    • Additionally, perform a crowded field analysis by including total compound counts, the percentage of marks from different owners, and a determination of whether the field is crowded.
    • Return your answer in JSON format with keys "components" and "crowded_field."
    
    YOUR RESPONSE MUST BE IN JSON FORMAT:
    {
      "identified_coordinated_classes": [LIST OF RELATED CLASS NUMBERS],
      "coordinated_classes_explanation": "[EXPLANATION OF WHY THESE CLASSES ARE RELATED TO THE PROPOSED TRADEMARK]",
      "components": [
        {
          "component": "[COMPONENT NAME]",
          "marks": [
            {
              "mark": "[TRADEMARK NAME]",
              "owner": "[OWNER NAME]",
              "goods_services": "[GOODS/SERVICES]",
              "status": "[LIVE/DEAD]",
              "class": "[CLASS]",
              "class_match": true|false,
              "goods_services_match": true|false
            }
          ],
          "distinctiveness": "[GENERIC|DESCRIPTIVE|SUGGESTIVE|ARBITRARY|FANCIFUL]"
        }
      ],
      "crowded_field": {
        "total_hits": [NUMBER],
        "distinct_owner_percentage": [PERCENTAGE],
        "is_crowded": true|false,
        "explanation": "[DETAILED EXPLANATION OF FINDINGS, INCLUDING REDUCED RISK IF is_crowded=true]"
      }
    }
"""  
  
    user_message = f"""
    Proposed Trademark: {mark}
    Class: {class_number}
    Goods/Services: {goods_services}
    
    Trademark Conflicts:
    {json.dumps(relevant_conflicts, indent=2)}
    
    Analyze ONLY Section II: Component Analysis.
    
    IMPORTANT REMINDERS:
    - Include exact counts and percentages for all statistics
    - For Crowded Field Analysis:
      1. Show the total number of compound mark hits
      2. Calculate percentage of marks with different owners
      3. If >50% have different owners, set is_crowded=true and mention decreased risk
    - For Class Match (True/False):
      * First, identify all coordinated classes related to the proposed goods/services "{goods_services}"
      * Mark True if the mark's class exactly matches the proposed class "{class_number}"
      * ALSO mark True if the mark's class is in a coordinated or related class grouping you identified
    - For Goods & Services Match (True/False), compare the mark's goods/services to the proposed goods/services "{goods_services}"
    - Always include the FULL goods/services description in your output, not just the class number
"""  
  
    try:  
        response = client.chat.completions.create(  
            model="gpt-4o",  
            messages=[  
                {"role": "system", "content": system_prompt},  
                {"role": "user", "content": user_message}  
            ],  
            temperature=0.0,  
        )  
  
        if response.choices and len(response.choices) > 0:  
            content = response.choices[0].message.content  
  
            # Extract JSON data  
            json_match = re.search(r'```json\s*(.*?)\s*```|({[\s\S]*})', content, re.DOTALL)  
            if json_match:  
                json_str = json_match.group(1) or json_match.group(2)  
                try:  
                    raw_results = json.loads(json_str)
                    return raw_results
                except json.JSONDecodeError:  
                    return {
                        "identified_coordinated_classes": [],
                        "coordinated_classes_explanation": "Unable to identify coordinated classes",
                        "components": [],  
                        "crowded_field": {  
                            "total_hits": 0,
                            "distinct_owner_percentage": 0,
                            "is_crowded": False,
                            "explanation": "Unable to determine crowded field status."  
                        }  
                    }  
            else:  
                return {
                    "identified_coordinated_classes": [],
                    "coordinated_classes_explanation": "Unable to identify coordinated classes",
                    "components": [],  
                    "crowded_field": {  
                        "total_hits": 0,
                        "distinct_owner_percentage": 0,
                        "is_crowded": False,
                        "explanation": "Unable to determine crowded field status."  
                    }  
                }  
        else:  
            return {
                "identified_coordinated_classes": [],
                "coordinated_classes_explanation": "Unable to identify coordinated classes",
                "components": [],  
                "crowded_field": {  
                    "total_hits": 0,
                    "distinct_owner_percentage": 0,
                    "is_crowded": False,
                    "explanation": "Unable to determine crowded field status."  
                }  
            }  
    except Exception as e:  
        print(f"Error in section_two_analysis: {str(e)}")  
        return {
            "identified_coordinated_classes": [],
            "coordinated_classes_explanation": "Error occurred during analysis",
            "components": [],  
            "crowded_field": {  
                "total_hits": 0,
                "distinct_owner_percentage": 0,
                "is_crowded": False,
                "explanation": "Error occurred during analysis"  
            }  
        }


def section_three_analysis(mark, class_number, goods_services, section_one_results, section_two_results):
    """
    Perform Section III: Risk Assessment and Summary
    
    Args:
        mark: The proposed trademark
        class_number: The class of the proposed trademark
        goods_services: The goods and services of the proposed trademark
        section_one_results: Results from Section I
        section_two_results: Results from Section II
        
    Returns:
        A structured risk assessment and summary
    """
    client = get_azure_client()
    
    system_prompt = """
    You are a trademark expert attorney specializing in trademark opinion writing.
    
    Please analyze the results from Sections I and II to create Section III: Risk Assessment and Summary. Your analysis should address:

    Likelihood of Confusion – Evaluate the potential for confusion between the proposed trademark and conflicting marks, including the impact of coordinated class conflicts.
    Descriptiveness – Assess whether the proposed trademark's goods/services are descriptive compared to the conflicts.
    Aggressive Enforcement and Litigious Behavior – Identify any patterns of aggressive enforcement among the owners of the conflicting marks.
    Overall Risk – Provide a risk rating (HIGH, MEDIUM-HIGH, MEDIUM, MEDIUM-LOW, LOW) along with an explanation.

    IMPORTANT:
    • In your risk analysis, consider the fact that a conflict may come from a coordinated or related class. For example, if a conflict is registered under a class different from 20 but falls within a related grouping (such as a class that frequently aligns with furniture or home furnishings), mention this and incorporate its impact on risk.
    • Also, include metrics from crowded field analysis to determine if overlapping conflicting marks reduce the overall risk.
    • Return your response in JSON format with keys: likelihood_of_confusion, descriptiveness, aggressive_enforcement, and overall_risk.
    
    YOUR RESPONSE MUST BE IN JSON FORMAT:
    {
      "likelihood_of_confusion": [
        "[KEY POINT ABOUT LIKELIHOOD OF CONFUSION]",
        "[ADDITIONAL POINT ABOUT LIKELIHOOD OF CONFUSION]"
      ],
      "descriptiveness": [
        "[KEY POINT ABOUT DESCRIPTIVENESS]"
      ],
      "aggressive_enforcement": {
        "owners": [
          {
            "name": "[OWNER NAME]",
            "enforcement_patterns": [
              "[PATTERN 1]",
              "[PATTERN 2]"
            ]
          }
        ],
        "enforcement_landscape": [
          "[KEY POINT ABOUT ENFORCEMENT LANDSCAPE]",
          "[ADDITIONAL POINT ABOUT ENFORCEMENT LANDSCAPE]"
        ]
      },
      "overall_risk": {
        "level": "[HIGH|MEDIUM-HIGH|MEDIUM|MEDIUM-LOW|LOW]",
        "explanation": "[EXPLANATION OF RISK LEVEL WITH FOCUS ON CROWDED FIELD]",
        "crowded_field_percentage": [PERCENTAGE]
      }
    }
    """
    
    user_message = f"""
    Proposed Trademark: {mark}
    Class: {class_number}
    Goods and Services: {goods_services}
    
    Section I Results:
    {json.dumps(section_one_results, indent=2)}
    
    Section II Results:
    {json.dumps(section_two_results, indent=2)}
    
    Create Section III: Risk Assessment and Summary.
    
    IMPORTANT REMINDERS:
    - Focus the risk discussion on crowded field analysis
    - Include the percentage of overlapping marks from crowded field analysis
    - Do NOT include recommendations
    - If the risk is Medium-High and a crowded field is identified, reduce it to Medium-Low
    - For aggressive enforcement analysis, examine the owners of similar marks and identify any known for litigious behavior
    - Specifically analyze coordinated class conflicts - marks in related class groupings may present significant risk even if they're not in the exact same class
    """
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message}
            ],
            temperature=0.0,
        )
        
        if response.choices and len(response.choices) > 0:
            content = response.choices[0].message.content
            
            # Extract JSON data
            json_match = re.search(r'```json\s*(.*?)\s*```|({[\s\S]*})', content, re.DOTALL)
            if json_match:
                json_str = json_match.group(1) or json_match.group(2)
                try:
                    return json.loads(json_str)
                except json.JSONDecodeError:
                    return {
                        "likelihood_of_confusion": ["Unable to determine likelihood of confusion."],
                        "descriptiveness": ["Unable to determine descriptiveness."],
                        "aggressive_enforcement": {
                            "owners": [],
                            "enforcement_landscape": ["Unable to determine enforcement patterns."]
                        },
                        "overall_risk": {
                            "level": "MEDIUM",
                            "explanation": "Unable to determine precise risk level.",
                            "crowded_field_percentage": 0
                        }
                    }
            else:
                return {
                    "likelihood_of_confusion": ["Unable to determine likelihood of confusion."],
                    "descriptiveness": ["Unable to determine descriptiveness."],
                    "aggressive_enforcement": {
                        "owners": [],
                        "enforcement_landscape": ["Unable to determine enforcement patterns."]
                    },
                    "overall_risk": {
                        "level": "MEDIUM",
                        "explanation": "Unable to determine precise risk level.",
                        "crowded_field_percentage": 0
                    }
                }
        else:
            return {
                "likelihood_of_confusion": ["Unable to determine likelihood of confusion."],
                "descriptiveness": ["Unable to determine descriptiveness."],
                "aggressive_enforcement": {
                    "owners": [],
                    "enforcement_landscape": ["Unable to determine enforcement patterns."]
                },
                "overall_risk": {
                    "level": "MEDIUM",
                    "explanation": "Unable to determine precise risk level.",
                    "crowded_field_percentage": 0
                }
            }
    except Exception as e:
        print(f"Error in section_three_analysis: {str(e)}")
        return {
            "likelihood_of_confusion": ["Unable to determine likelihood of confusion."],
            "descriptiveness": ["Unable to determine descriptiveness."],
            "aggressive_enforcement": {
                "owners": [],
                "enforcement_landscape": ["Unable to determine enforcement patterns."]
            },
            "overall_risk": {
                "level": "MEDIUM",
                "explanation": "Unable to determine precise risk level.",
                "crowded_field_percentage": 0
            }
        }

def generate_trademark_opinion(conflicts_array, proposed_name, proposed_class, proposed_goods_services):
    """
    Generate a comprehensive trademark opinion by running the entire analysis process.
    
    Args:
        conflicts_array: List of potential trademark conflicts
        proposed_name: Name of the proposed trademark
        proposed_class: Class of the proposed trademark
        proposed_goods_services: Goods and services description
        
    Returns:
        A comprehensive trademark opinion
    """
    # Pre-filter trademarks to get the excluded count
    relevant_conflicts, excluded_count = validate_trademark_relevance(conflicts_array, proposed_goods_services)
    
    print("Performing Section I: Comprehensive Trademark Hit Analysis...")
    section_one_results = section_one_analysis(proposed_name, proposed_class, proposed_goods_services, relevant_conflicts)
    
    print("Performing Section II: Component Analysis...")
    section_two_results = section_two_analysis(proposed_name, proposed_class, proposed_goods_services, relevant_conflicts)
    
    print("Performing Section III: Risk Assessment and Summary...")
    section_three_results = section_three_analysis(proposed_name, proposed_class, proposed_goods_services, section_one_results, section_two_results)
    
    # Create a comprehensive opinion structure
    opinion_structure = {
        "proposed_name": proposed_name,
        "proposed_class": proposed_class,
        "proposed_goods_services": proposed_goods_services,
        "excluded_count": excluded_count,
        "section_one": section_one_results,
        "section_two": section_two_results,
        "section_three": section_three_results
    }
    
    # Format the opinion in a structured way
    comprehensive_opinion = f"""
    REFINED TRADEMARK OPINION: {proposed_name}
    Class: {proposed_class}
    Goods and Services: {proposed_goods_services}

    Section I: Comprehensive Trademark Hit Analysis
    
    (a) Identical Marks:
    {json.dumps(section_one_results.get('identical_marks', []), indent=2)}
    
    (b) One Letter and Two Letter Analysis:
    {json.dumps({
        'one_letter_marks': section_one_results.get('one_letter_marks', []),
        'two_letter_marks': section_one_results.get('two_letter_marks', [])
    }, indent=2)}
    
    (c) Phonetically, Semantically & Functionally Similar Analysis:
    {json.dumps(section_one_results.get('similar_marks', []), indent=2)}
    
    (d) Crowded Field Analysis:
    {json.dumps(section_one_results.get('crowded_field', {}), indent=2)}

    Section II: Component Analysis
    
    (a) Component Analysis:
    {json.dumps(section_two_results.get('components', []), indent=2)}
    
    (b) Crowded Field Analysis:
    {json.dumps(section_two_results.get('crowded_field', {}), indent=2)}

    Section III: Risk Assessment and Summary
    
    Likelihood of Confusion:
    {json.dumps(section_three_results.get('likelihood_of_confusion', []), indent=2)}
    
    Descriptiveness:
    {json.dumps(section_three_results.get('descriptiveness', []), indent=2)}
    
    Overall Risk Level:
    {json.dumps(section_three_results.get('overall_risk', {}), indent=2)}
    
    Note: {excluded_count} trademarks with unrelated goods/services were excluded from this analysis.
    """
    
    # Clean and format the final opinion
    print("Cleaning and formatting the final opinion...")
    formatted_opinion = clean_and_format_opinion(comprehensive_opinion, opinion_structure)
    
    return formatted_opinion


# Example usage function
def run_trademark_analysis(proposed_name, proposed_class, proposed_goods_services, conflicts_data):
    """
    Run a complete trademark analysis with proper error handling.
    
    Args:
        proposed_name: Name of the proposed trademark
        proposed_class: Class of the proposed trademark
        proposed_goods_services: Goods and services of the proposed trademark
        conflicts_data: Array of potential conflict trademarks
        
    Returns:
        A comprehensive trademark opinion
    """
    try:
        if not proposed_name or not proposed_class or not proposed_goods_services:
            return "Error: Missing required trademark information."
            
        if not conflicts_data:
            return "Error: No conflict data provided for analysis."
            
        opinion = generate_trademark_opinion(conflicts_data, proposed_name, proposed_class, proposed_goods_services)
        return opinion
        
    except Exception as e:
        return f"Error running trademark analysis: {str(e)}"

# TAMIL CODE END'S HERE ---------------------------------------------------------------------------------------------------------------------------

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def export_trademark_opinion_to_word(opinion_output):
    """
    Export trademark opinion to Word document with proper formatting and table support
    """
    document = Document()
    
    # Track if we're currently building a table
    current_table = None
    header_row = []
    
    # Parse and handle different sections
    lines = opinion_output.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Handle table headers
        if line.startswith('|') and ('Trademark' in line or 'Cited Term' in line or 'Component' in line):
            # Extract header cells
            header_cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            
            # Get the separator line that defines the table structure
            if i + 1 < len(lines) and '---' in lines[i+1]:
                i += 1  # Skip the separator line
            
            # Create a new table
            current_table = document.add_table(rows=1, cols=len(header_cells))
            current_table.style = 'Table Grid'
            
            # Add header row
            for j, cell in enumerate(header_cells):
                if j < len(current_table.rows[0].cells):
                    current_table.rows[0].cells[j].text = cell
            
            header_row = header_cells
            
        # Handle table data rows
        elif line.startswith('|') and current_table is not None:
            cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            
            if cells:  # Only proceed if we have cells
                # Add a new row
                new_row = current_table.add_row()
                
                # Fill the row with data, ensuring we don't exceed the number of columns
                for j, cell in enumerate(cells):
                    if j < len(new_row.cells):
                        new_row.cells[j].text = cell
        
        # Handle regular paragraphs (section headers, etc.)
        elif line and not line.startswith('|'):
            # If line is a section header (ends with :)
            if line.endswith(':') or line.startswith('Section'):
                p = document.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                p.space_after = Pt(12)
            else:
                # Regular paragraph
                document.add_paragraph(line)
                
            # Reset table tracking when we hit a new paragraph
            if line.strip() and current_table is not None:
                current_table = None
        
        i += 1
    
    # Save the document
    filename = "Trademark_Opinion.docx"
    document.save(filename)
    return filename 

# -------------------------------------------------------------

from typing import List  
import fitz  # PyMuPDF  
from PIL import Image  
import io  
  
  
# def Web_CommonLaw_Overview_List(document: str, start_page: int, pdf_document: fitz.Document) -> List[int]:  
#     """  
#     Extract the page numbers for the 'Web Common Law Overview List' section.  
#     """  
#     pages_with_overview = []  
#     for i in range(start_page, min(start_page + 2, pdf_document.page_count)):  
#         page = pdf_document.load_page(i)  
#         page_text = page.get_text()  
#         if "Record Nr." in page_text:  # Check for "Record Nr." in the text  
#             pages_with_overview.append(i + 1)  # Use 1-based indexing for page numbers  
#     return pages_with_overview  
  
  
# def convert_pages_to_pil_images(pdf_document: fitz.Document, page_numbers: List[int]) -> List[Image.Image]:  
#     """  
#     Convert the specified pages of the PDF to PIL images and return them as a list of PIL Image objects.  
#     """  
#     images = []  
#     for page_num in page_numbers:  
#         page = pdf_document.load_page(page_num - 1)  # Convert 1-based index to 0-based  
#         pix = page.get_pixmap()  # Render the page to a pixmap  
#         img = Image.open(io.BytesIO(pix.tobytes("png")))  # Convert pixmap to PIL Image  
#         images.append(img)  # Add the PIL Image object to the list  
#     return images  
  
  
# def web_law_page(document_path: str) -> List[Image.Image]:  
#     """  
#     Return PIL Image objects of the pages where either:  
#     1. "Web Common Law Summary Page:" appears, or  
#     2. Both "Web Common Law Overview List" and "Record Nr." appear.  
#     """  
#     matching_pages = []  # List to store matching page numbers  
  
#     with fitz.open(document_path) as pdf_document:  
#         for page_num in range(pdf_document.page_count):  
#             page = pdf_document.load_page(page_num)  
#             page_text = page.get_text()  
#             print(page_text)  
              
#             # Check for "Web Common Law Summary Page:"  
#             if "Web Common Law Page:" in page_text:  
#                 matching_pages.append(page_num + 1)  
  
  
#             # Check for "Web Common Law Overview List" and "Record Nr."  
#             if "WCL-" in page_text:  
#                 matching_pages.append(page_num + 1)  
#             # if "Web Common Law Overview List" in page_text and "Record Nr." in page_text:  
#             #     overview_pages = Web_CommonLaw_Overview_List(  
#             #         page_text, page_num, pdf_document  
#             #     )  
#             #     matching_pages.extend(overview_pages)  
  
  
#         # Remove duplicates and sort the page numbers  
#         matching_pages = sorted(set(matching_pages))  
  
#         # Convert matching pages to PIL images  
#         images = convert_pages_to_pil_images(pdf_document, matching_pages)  
  
#     return images  
                
# # ---- extraction logic

# import io  
# import base64  
# import cv2  
# import json  
# import requests  
# import os
# from PIL import Image  
# from typing import List  
# import numpy as np
  
# # Function to encode images using OpenCV  
# def encode_image(image: Image.Image) -> str:  
#     """  
#     Encode a PIL Image as Base64 string using OpenCV.  
#     """  
#     # Convert PIL Image to numpy array for OpenCV  
#     image_np = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)  
#     buffered = cv2.imencode(".jpg", image_np)[1]  
#     return base64.b64encode(buffered).decode("utf-8")  
  
  
# # Function to process a single image and get the response from LLM  
# def process_single_image(image: Image.Image, proposed_name: str) -> dict:  
#     """  
#     Process a single image by sending it to Azure OpenAI API.  
#     Cited term: Check for {proposed_name} in the image.
#     """        
#     azure_endpoint = os.getenv("AZURE_ENDPOINT")
#     api_key = os.getenv("AZURE_API_KEY")
#     model="gpt-4o"  

#     # Encode the image into Base64 using OpenCV  
#     base64_image = encode_image(image)  
  
#     # Prepare the prompt for the LLM  
#     prompt = f"""Extract the following details from the given image: Cited term, Owner name, Goods & services.\n\n
    
#                 Cited Term:\n
#                 - This is the snippet in the product/site text that *fully or partially matches* the physically highlighted or searched trademark name: {proposed_name}.
#                 - You must prioritize any match that closely resembles '{proposed_name}' — e.g., 'ColorGrip', 'COLORGRIP', 'Color self Grip' , 'Grip Colour', 'color-grip', 'Grip' , or minor variations in spacing/punctuation.

#                 Owner Name (Brand):\n
#                 - Identify the name of the individual or entity that owns or manufactures the product.
#                 - Look for indicators like "Owner:," "Brand:," "by:," or "Manufacturer:."
#                 - If none are found, return "Not specified."
                
#                 Goods & Services:\n
#                 - Extract the core goods and services associated with the trademark or product.  
#                 - Provide relevant detail (e.g., "permanent hair color," "nail care polish," "hair accessories," or "hair styling tools").
    
#                 Return output only in the exact below-mentioned format:  
#                 Example output format:  
#                     Cited_term: ColourGrip,\n  
#                     Owner_name: Matrix, \n 
#                     Goods_&_services: Hair color products,\n    
# """
  
#     # Prepare the API payload  
#     data = {  
#         "model": model,  
#         "messages": [  
#             {  
#                 "role": "system",  
#                 "content": "You are a helpful assistant for extracting Meta Data based on the given Images [Note: Only return the required extracted data in the exact format mentioned].",  
#             },  
#             {  
#                 "role": "user",  
#                 "content": [  
#                     {"type": "text", "text": prompt},  
#                     {  
#                         "type": "image_url",  
#                         "image_url": {  
#                             "url": f"data:image/png;base64,{base64_image}"  
#                         },  
#                     },  
#                 ],  
#             },  
#         ],  
#         "max_tokens": 200,  
#         "temperature": 0,  
#     }  
  
#     # Send the API request  
#     headers = {"Content-Type": "application/json", "api-key": api_key}  
#     response = requests.post(  
#         f"{azure_endpoint}/openai/deployments/{model}/chat/completions?api-version=2024-10-01-preview",  
#         headers=headers,  
#         data=json.dumps(data),  
#     )  
  
#     # Parse the response  
#     if response.status_code == 200:  
#         extracted_data = response.json()["choices"][0]["message"]["content"]  
#     else:  
#         extracted_data = "Failed to extract data"    
#     # Return the extracted data  
#     return {extracted_data.strip()}  
  
  
# # Function to process all images one by one  
# def extract_web_common_law(page_images: List[Image.Image], proposed_name: str) -> List[dict]:  
#     """  
#     Send images one by one to Azure OpenAI GPT models,  
#     and collect the responses into a single array.  
#     """    
#     # Process each image and collect the results  
#     results = []  
#     for idx, image in enumerate(page_images):  
#         result = process_single_image(image, proposed_name)  
#         results.append(result)  
  
#     # Return the collected results as a single array  
#     return results  

# def analyze_web_common_law(extracted_data: List[str], proposed_name: str) -> str:
#     """
#     Comprehensive analysis of web common law trademark data through three specialized stages.
#     Returns a professional opinion formatted according to legal standards.
#     """
#     # Stage 1: Cited Term Analysis
#     cited_term_analysis = perform_cited_term_analysis(extracted_data, proposed_name)
    
#     # Stage 2: Component Analysis
#     component_analysis = perform_component_analysis(extracted_data, proposed_name)
    
#     # Stage 3: Final Risk Assessment
#     risk_assessment = perform_risk_assessment(cited_term_analysis, component_analysis, proposed_name)
    
#     # Combine all sections into final report
#     final_report = f"""
# WEB COMMON LAW OPINION: {proposed_name}

# {cited_term_analysis}

# {component_analysis}

# {risk_assessment}
# """
#     return final_report

# def perform_cited_term_analysis(extracted_data: List[str], proposed_name: str) -> str:
#     """
#     Perform Section IV: Comprehensive Cited Term Analysis
#     """
#     azure_endpoint = os.getenv("AZURE_ENDPOINT")
#     api_key = os.getenv("AZURE_API_KEY")
#     model = "gpt-4o"

#     extracted_text = "\n".join([str(item) for item in extracted_data])
    
#     prompt = f"""You are a trademark attorney analyzing web common law trademark data.
# Perform Section IV analysis (Comprehensive Cited Term Analysis) with these subsections:

# 1. Identical Cited Terms
# 2. One Letter and Two Letter Differences
# 3. Phonetically/Semantically/Functionally Similar Terms

# Analyze this web common law data against proposed trademark: {proposed_name}

# Extracted Data:
# {extracted_text}

# Perform comprehensive analysis:
# 1. Check for identical cited terms
# 2. Analyze one/two letter differences
# 3. Identify similar terms (phonetic/semantic/functional)
# 4. For each, determine if goods/services are similar

# Return results in EXACTLY this format:

# Section IV: Comprehensive Cited Term Analysis
# (a) Identical Cited Terms:
# | Cited Term | Owner | Goods & Services | Goods & Services Match |
# |------------|--------|------------------|------------------------|
# | [Term 1] | [Owner] | [Goods/Services] | [True/False] |

# (b) One Letter and Two Letter Analysis:
# | Cited Term | Owner | Goods & Services | Difference Type | Goods & Services Match |
# |------------|--------|------------------|----------------|------------------------|
# | [Term 1] | [Owner] | [Goods/Services] | [One/Two Letter] | [True/False] |

# (c) Phonetically, Semantically & Functionally Similar Analysis:
# | Cited Term | Owner | Goods & Services | Similarity Type | Goods & Services Match |
# |------------|--------|------------------|-----------------|------------------------|
# | [Term 1] | [Owner] | [Goods/Services] | [Phonetic/Semantic/Functional] | [True/False] |

# Evaluation Guidelines:
# - Goods/services match if they overlap with proposed trademark's intended use
# - One letter difference = exactly one character changed/added/removed
# - Two letter difference = exactly two characters changed/added/removed
# - Phonetic similarity = sounds similar when spoken
# - Semantic similarity = similar meaning
# - Functional similarity = similar purpose/use
# - State "None" when no results are found
# - Filter out rows where both match criteria are False
# - Always include complete goods/services text
# """

#     data = {
#         "model": model,
#         "messages": [
#             {
#                 "role": "system",
#                 "content": "You are a trademark attorney specializing in comprehensive trademark analysis. Provide precise, professional analysis in the exact requested format.",
#             },
#             {
#                 "role": "user",
#                 "content": prompt,
#             },
#         ],
#         "max_tokens": 2000,
#         "temperature": 0.1,
#     }

#     headers = {"Content-Type": "application/json", "api-key": api_key}
#     response = requests.post(
#         f"{azure_endpoint}/openai/deployments/{model}/chat/completions?api-version=2024-10-01-preview",
#         headers=headers,
#         data=json.dumps(data),
#     )

#     if response.status_code == 200:
#         return response.json()["choices"][0]["message"]["content"]
#     return "Failed to generate cited term analysis"

# def perform_component_analysis(extracted_data: List[str], proposed_name: str) -> str:
#     """
#     Perform Section V: Component Analysis and Crowded Field Assessment
#     """
#     azure_endpoint = os.getenv("AZURE_ENDPOINT")
#     api_key = os.getenv("AZURE_API_KEY")
#     model = "gpt-4o"

#     extracted_text = "\n".join([str(item) for item in extracted_data])
    
#     prompt = f"""You are a trademark attorney analyzing web common law components.
# Perform Section V analysis (Component Analysis) with these subsections:

# 1. Component Breakdown
# 2. Crowded Field Analysis

# Analyze this web common law data against proposed trademark: {proposed_name}

# Extracted Data:
# {extracted_text}

# Perform component analysis:
# 1. Break proposed term into meaningful components
# 2. For each component, find other terms using that component
# 3. Perform crowded field analysis
# 4. Determine goods/services matches

# Return results in EXACTLY this format:

# Section V: Component Analysis
# (a) Component Analysis:

# Component 1: [First Component]
# | Cited Term | Owner | Goods & Services | Goods & Services Match |
# |-----------|--------|------------------|------------------------|
# | [Term 1] | [Owner] | [Goods/Services] | [True/False] |

# (b) Crowded Field Analysis:
# - **Total component hits found**: [NUMBER]
# - **Terms with different owners**: [NUMBER] ([PERCENTAGE]%)
# - **Crowded Field Status**: [YES/NO]
# - **Analysis**: 
#   [DETAILED EXPLANATION OF FINDINGS INCLUDING RISK IMPLICATIONS IF FIELD IS CROWDED]

# IMPORTANT:
# 1. Break cited term into meaningful components
# 2. For each component, find other terms using that component
# 3. Include FULL goods/services descriptions
# 4. For crowded field:
#    - Calculate percentage of distinct owners
#    - Field is crowded if >50% different owners
#    - Include detailed explanation
# 5. Assess distinctiveness for each component

# Additional Instructions:
# - Goods/services match if they overlap with proposed trademark's intended use
# - For crowded field, calculate:
#   * Total component hits
#   * Percentage with different owners
# - Distinctiveness levels:
#   * Generic: Common term for the goods/services
#   * Descriptive: Describes characteristic/quality
#   * Suggestive: Suggests qualities (requires imagination)
#   * Arbitrary: Common word unrelated to goods/services
#   * Fanciful: Invented word
# - State "None" when no results are found
# - Filter out rows where both match criteria are False
# - Always include complete goods/services text
# """

#     data = {
#         "model": model,
#         "messages": [
#             {
#                 "role": "system",
#                 "content": "You are a trademark attorney specializing in component and crowded field analysis. Provide precise, professional analysis in the exact requested format.",
#             },
#             {
#                 "role": "user",
#                 "content": prompt,
#             },
#         ],
#         "max_tokens": 2000,
#         "temperature": 0.1,
#     }

#     headers = {"Content-Type": "application/json", "api-key": api_key}
#     response = requests.post(
#         f"{azure_endpoint}/openai/deployments/{model}/chat/completions?api-version=2024-10-01-preview",
#         headers=headers,
#         data=json.dumps(data),
#     )

#     if response.status_code == 200:
#         return response.json()["choices"][0]["message"]["content"]
#     return "Failed to generate component analysis"

# def perform_risk_assessment(cited_term_analysis: str, component_analysis: str, proposed_name: str) -> str:
#     """
#     Perform Section VI: Final Risk Assessment combining all findings
#     """
#     azure_endpoint = os.getenv("AZURE_ENDPOINT")
#     api_key = os.getenv("AZURE_API_KEY")
#     model = "gpt-4o"

#     prompt = f"""You are a senior trademark attorney preparing a final risk assessment.
# Combine these analysis sections into a comprehensive risk assessment for: {proposed_name}

# Cited Term Analysis:
# {cited_term_analysis}

# Component Analysis:
# {component_analysis}

# Prepare Section VI: Web Common Law Risk Assessment with these subsections:

# 1. Market Presence
# 2. Enforcement Patterns
# 3. Risk Category for Use
# 4. Combined Risk Assessment

# Return results in EXACTLY this format:

# Section VI: Web Common Law Risk Assessment

# Market Presence:
# - [KEY POINT ABOUT MARKET PRESENCE]

# Enforcement Patterns:
# - **Known Aggressive Owners**:
#   * [Owner 1]: [Enforcement patterns]

# Risk Category for Use:
# - **[USE RISK LEVEL: HIGH/MEDIUM/LOW]**
# - [EXPLANATION OF USE RISK LEVEL]

# III. COMBINED RISK ASSESSMENT

# Overall Risk Category:
# - **[OVERALL RISK LEVEL: HIGH/MEDIUM-HIGH/MEDIUM/MEDIUM-LOW/LOW]**
# - [EXPLANATION INCORPORATING BOTH TRADEMARK AND WEB COMMON LAW FINDINGS]

# Guidelines:
# 1. Base assessment strictly on the provided analysis
# 2. Do not introduce new findings not in the analysis
# 3. Maintain professional, legal tone
# 4. Be specific about risk factors
# 5. Highlight any particularly concerning findings
# """

#     data = {
#         "model": model,
#         "messages": [
#             {
#                 "role": "system",
#                 "content": "You are a senior trademark attorney specializing in risk assessment. Provide precise, professional analysis in the exact requested format.",
#             },
#             {
#                 "role": "user",
#                 "content": prompt,
#             },
#         ],
#         "max_tokens": 1500,
#         "temperature": 0.1,
#     }

#     headers = {"Content-Type": "application/json", "api-key": api_key}
#     response = requests.post(
#         f"{azure_endpoint}/openai/deployments/{model}/chat/completions?api-version=2024-10-01-preview",
#         headers=headers,
#         data=json.dumps(data),
#     )

#     if response.status_code == 200:
#         return response.json()["choices"][0]["message"]["content"]
#     return "Failed to generate risk assessment"

# -------------------

# Streamlit App
st.title("Trademark Document Parser Version 6.9")

# File upload
uploaded_files = st.sidebar.file_uploader(
    "Choose PDF files", type="pdf", accept_multiple_files=True
)

if uploaded_files:
    if st.sidebar.button("Check Conflicts", key="check_conflicts"):
        total_files = len(uploaded_files)
        progress_bar = st.progress(0)
        # progress_label.text(f"Progress: 0%")  --- Needed to set

        for i, uploaded_file in enumerate(uploaded_files):
            # Save uploaded file to a temporary file path
            temp_file_path = f"temp_{uploaded_file.name}"
            with open(temp_file_path, "wb") as f:
                f.write(uploaded_file.read())

            start_time = time.time()

            sp = True
            proposed_trademark_details = extract_proposed_trademark_details(
                temp_file_path
            )

            if proposed_trademark_details:
                proposed_name = proposed_trademark_details.get(
                    "proposed_trademark_name", "N"
                )
                proposed_class = proposed_trademark_details.get(
                    "proposed_nice_classes_number"
                )
                proposed_goods_services = proposed_trademark_details.get(
                    "proposed_goods_services", "N"
                )
                if proposed_goods_services != "N":
                    with st.expander(
                        f"Proposed Trademark Details for {uploaded_file.name}"
                    ):
                        st.write(f"Proposed Trademark name: {proposed_name}")
                        st.write(f"Proposed class-number: {proposed_class}")
                        st.write(
                            f"Proposed Goods & Services: {proposed_goods_services}"
                        )
                    class_list = list_conversion(proposed_class)
                else:
                    st.write(
                        "______________________________________________________________________________________________________________________________"
                    )
                    st.write(
                        f"Sorry, unable to generate report due to insufficient information about goods & services in the original trademark report : {uploaded_file.name}"
                    )
                    st.write(
                        "______________________________________________________________________________________________________________________________"
                    )
                    sp = False
            else:

                proposed_trademark_details = extract_proposed_trademark_details2(
                    temp_file_path
                )

                if proposed_trademark_details:
                    proposed_name = proposed_trademark_details.get(
                        "proposed_trademark_name", "N"
                    )
                    proposed_class = proposed_trademark_details.get(
                        "proposed_nice_classes_number"
                    )
                    proposed_goods_services = proposed_trademark_details.get(
                        "proposed_goods_services", "N"
                    )
                    if proposed_goods_services != "N":
                        with st.expander(
                            f"Proposed Trademark Details for {uploaded_file.name}"
                        ):
                            st.write(f"Proposed Trademark name: {proposed_name}")
                            st.write(f"Proposed class-number: {proposed_class}")
                            st.write(
                                f"Proposed Goods & Services: {proposed_goods_services}"
                            )
                        class_list = list_conversion(proposed_class)
                    else:
                        st.write(
                            "______________________________________________________________________________________________________________________________"
                        )
                        st.write(
                            f"Sorry, unable to generate report due to insufficient information about goods & services in the original trademark report : {uploaded_file.name}"
                        )
                        st.write(
                            "______________________________________________________________________________________________________________________________"
                        )
                        sp = False
                else:
                    st.error(
                        f"Unable to extract Proposed Trademark Details for {uploaded_file.name}"
                    )
                    sp = False
                    continue

            if sp:
                progress_bar.progress(25)
                # Initialize AzureChatOpenAI

                # s_time = time.time()

                existing_trademarks = parse_trademark_details(temp_file_path)
                st.write(len(existing_trademarks))
                # for i in range(25,46):
                #     progress_bar.progress(i)


# # PRAVEEN WEB COMMON LAW CODE START'S HERE-------------------------------------------------------------------------------------------------------------------------

#                 # Updated usage in your Streamlit code would look like:
#                 # !!! Function used extract the web common law pages into images
#                 full_web_common_law = web_law_page(temp_file_path)                

#                 progress_bar.progress(50)
#                 st.success(
#                     f"Existing Trademarks Data Extracted Successfully for {uploaded_file.name}!"
#                 )

#                 # !!! Function used extract the web common law details from the images using LLM 
#                 extracted_web_law = extract_web_common_law(full_web_common_law, proposed_name)  

#                 # New comprehensive analysis
#                 analysis_result = analyze_web_common_law(extracted_web_law, proposed_name)

#                 # Display results
#                 with st.expander("Extracted Web Common Law Data"):
#                     st.write(extracted_web_law)

#                 with st.expander("Trademark Legal Analysis"):
#                     st.markdown(analysis_result)  # Using markdown for better formatting

#                 # extracted_web_law ----- Web common law stored in this variable 

# PRAVEEN WEB COMMON LAW CODE END'S HERE-------------------------------------------------------------------------------------------------------------------------


                # e_time = time.time()
                # elap_time = e_time - s_time
                # elap_time = elap_time // 60
                # st.write(f"Time taken for extraction: {elap_time} mins")

                # e_time = time.time()
                # elap_time = e_time - s_time
                # st.write(f"Time taken: {elap_time} seconds")

                # Display extracted details

                nfiltered_list = []
                unsame_class_list = []

                # Iterate over each JSON element in trademark_name_list
                for json_element in existing_trademarks:
                    class_numbers = json_element["international_class_number"]
                    # Check if any of the class numbers are in class_list
                    if any(number in class_list for number in class_numbers):
                        nfiltered_list.append(json_element)
                    else:
                        unsame_class_list.append(json_element)

                existing_trademarks = nfiltered_list
                existing_trademarks_unsame = unsame_class_list


                for existing_trademark in existing_trademarks:
                    section_one_analysis(proposed_name, proposed_class, proposed_goods_services, existing_trademark)

                st.write("+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++")

                for existing_trademark in existing_trademarks_unsame:
                    if existing_trademark["international_class_number"] != []:
                        section_one_analysis(proposed_name, proposed_class, proposed_goods_services, existing_trademark)

                high_conflicts = []
                moderate_conflicts = []
                low_conflicts = []
                Name_Matchs = []
                no_conflicts = []

                lt = len(existing_trademarks)

                for existing_trademark in existing_trademarks:
                    conflict = compare_trademarks(
                        existing_trademark,
                        proposed_name,
                        proposed_class,
                        proposed_goods_services,
                    )
                    
                    if conflict is not None:
                        if conflict["conflict_grade"] == "High":
                            high_conflicts.append(conflict)
                        elif conflict["conflict_grade"] == "Moderate":
                            moderate_conflicts.append(conflict)
                        elif conflict["conflict_grade"] == "Low":
                            low_conflicts.append(conflict)
                        else:
                            no_conflicts.append(conflict)

                for existing_trademarks in existing_trademarks_unsame:
                    if existing_trademarks["international_class_number"] != []:
                        conflict = assess_conflict(
                            existing_trademarks,
                            proposed_name,
                            proposed_class,
                            proposed_goods_services,
                        )

                        if conflict["conflict_grade"] == "Name-Match":
                            # conflict_validation = compare_trademarks2(existing_trademarks, proposed_name, proposed_class, proposed_goods_services)
                            # if conflict_validation == "Name-Match":
                            Name_Matchs.append(conflict)
                        else:
                            print("Low")
                            # low_conflicts.append(conflict)

                st.sidebar.write("_________________________________________________")
                st.sidebar.subheader("\n\nConflict Grades : \n")
                st.sidebar.markdown(f"File: {proposed_name}")
                st.sidebar.markdown(
                    f"Total number of conflicts: {len(high_conflicts) + len(moderate_conflicts) + len(Name_Matchs) + len(low_conflicts)}"
                )
                st.sidebar.markdown(f"3 conditions satisfied:  {len(high_conflicts)}")
                st.sidebar.markdown(f"2 conditions satisfied:  {len(moderate_conflicts)}")
                st.sidebar.markdown(f"Name Match's Conflicts: {len(Name_Matchs)}")
                st.sidebar.markdown(f"1 condition satisfied: {len(low_conflicts)}")
                st.sidebar.write("_________________________________________________")

                document = Document()

                # Set page size to landscape  
                section = document.sections[0]  
                new_width, new_height = section.page_height, section.page_width  
                section.page_width = new_width  
                section.page_height = new_height  

                document.add_heading(
                    f"Trademark Conflict List for {proposed_name} (VERSION - 6.9) :"
                )

                document.add_heading("Dashboard :", level=2)
                # document.add_paragraph(f"\n\nTotal number of conflicts: {len(high_conflicts) + len(moderate_conflicts) + len(Name_Matchs) + len(low_conflicts)}\n- High Conflicts: {len(high_conflicts)}\n- Moderate Conflicts: {len(moderate_conflicts)}\n- Name Match's Conflicts: {len(Name_Matchs)}\n- Low Conflicts: {len(low_conflicts)}\n")

                # Updated Calculate the number of conflicts
                total_conflicts = (
                    len(high_conflicts)
                    + len(moderate_conflicts)
                    + len(Name_Matchs)
                    + len(low_conflicts)
                )

                # Create a table with 5 rows (including the header) and 2 columns
                table = document.add_table(rows=5, cols=2)

                # Set the table style and customize the borders
                table.style = "TableGrid"

                tbl = table._tbl
                tblBorders = OxmlElement("w:tblBorders")

                for border in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                    border_element = OxmlElement(f"w:{border}")
                    border_element.set(qn("w:val"), "single")
                    border_element.set(
                        qn("w:sz"), "4"
                    )  # This sets the border size; you can adjust it as needed
                    border_element.set(qn("w:space"), "0")
                    border_element.set(qn("w:color"), "000000")
                    tblBorders.append(border_element)

                tbl.append(tblBorders)

                # Fill the first column with labels
                labels = [
                    "Total number of conflicts:",
                    "- 3 conditions satisfied:",
                    "- 2 conditions satisfied:",
                    "- Name Match's Conflicts:",
                    "- 1 condition satisfied:",
                ]

                # Fill the second column with the conflict numbers
                values = [
                    total_conflicts,
                    len(high_conflicts),
                    len(moderate_conflicts),
                    len(Name_Matchs), 
                    len(low_conflicts),
                ]

                p = document.add_paragraph(" ")
                p.paragraph_format.line_spacing = Pt(18)
                p.paragraph_format.space_after = Pt(0)

                document.add_heading(
                    "Trademark Definitions: ", level=2
                )
                # p = document.add_paragraph(" ")
                # p.paragraph_format.line_spacing = Pt(18)
                p = document.add_paragraph("CONDITION 1: MARK: NAME-BASED SIMILARITY (comprised of Exact Match, Semantically Equivalent, Phonetically Equivalent, Primary position match)")
                p.paragraph_format.line_spacing = Pt(18)
                p.paragraph_format.space_after = Pt(0)
                p = document.add_paragraph("CONDITION 2: CLASS: CLASS OVERLAP")
                p.paragraph_format.line_spacing = Pt(18)
                p.paragraph_format.space_after = Pt(0)
                p = document.add_paragraph("CONDITION 3: GOODS/SERVICES: OVERLAPPING GOODS/SERVICES & TARGET MARKETS")
                p.paragraph_format.line_spacing = Pt(18)
                p.paragraph_format.space_after = Pt(0)
                p = document.add_paragraph("DIRECT HIT: Direct Name hit, regardless of the class")
                p.paragraph_format.line_spacing = Pt(18)
                p.paragraph_format.space_after = Pt(0)
                p = document.add_paragraph(" ")
                p.paragraph_format.line_spacing = Pt(18)
                p.paragraph_format.space_after = Pt(0)


                # Populate the table with the labels and values
                for i in range(5):
                    table.cell(i, 0).text = labels[i]
                    table.cell(i, 1).text = str(values[i])

                    # Set the font size to 10 for both cells
                    for cell in table.row_cells(i):
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(10)

                if len(high_conflicts) > 0:
                    document.add_heading("Trademarks with 3 conditions satisfied:", level=2)
                    # Create a pandas DataFrame from the JSON list
                    df_high = pd.DataFrame(high_conflicts)
                    df_high = df_high.drop(
                        columns=[
                            "Trademark name",
                            "Trademark class Number",
                            "Trademark registration number",
                            "Trademark serial number",
                            "Trademark design phrase",
                            "conflict_grade",
                            "reasoning",
                        ]
                    )
                    # Create a table in the Word document
                    table_high = document.add_table(
                        df_high.shape[0] + 1, df_high.shape[1]
                    )
                    # Set a predefined table style (with borders)
                    table_high.style = (
                        "TableGrid"  # This is a built-in style that includes borders
                    )
                    # Add the column names to the table
                    for i, column_name in enumerate(df_high.columns):
                        table_high.cell(0, i).text = column_name
                    # Add the data to the table
                    for i, row in df_high.iterrows():
                        for j, value in enumerate(row):
                            cell = table_high.cell(i + 1, j)
                            cell.text = str(value)
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)

                if len(moderate_conflicts) > 0:
                    document.add_heading("Trademarks with 2 conditions satisfied:", level=2)
                    # Create a pandas DataFrame from the JSON list
                    df_moderate = pd.DataFrame(moderate_conflicts)
                    df_moderate = df_moderate.drop(
                        columns=[
                            "Trademark name",
                            "Trademark class Number",
                            "Trademark registration number",
                            "Trademark serial number",
                            "Trademark design phrase",
                            "conflict_grade",
                            "reasoning",
                        ]
                    )
                    # Create a table in the Word document
                    table_moderate = document.add_table(
                        df_moderate.shape[0] + 1, df_moderate.shape[1]
                    )
                    # Set a predefined table style (with borders)
                    table_moderate.style = (
                        "TableGrid"  # This is a built-in style that includes borders
                    )
                    # Add the column names to the table
                    for i, column_name in enumerate(df_moderate.columns):
                        table_moderate.cell(0, i).text = column_name
                    # Add the data to the table
                    for i, row in df_moderate.iterrows():
                        for j, value in enumerate(row):
                            cell = table_moderate.cell(i + 1, j)
                            cell.text = str(value)
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)

                if len(Name_Matchs) > 0:
                    document.add_heading(
                        "Trademarks with Name Match's Conflicts:", level=2
                    )
                    # Create a pandas DataFrame from the JSON list
                    df_Name_Matchs = pd.DataFrame(Name_Matchs)
                    df_Name_Matchs = df_Name_Matchs.drop(
                        columns=[
                            "Trademark name",
                            "Trademark class Number",
                            "Trademark registration number",
                            "Trademark serial number",
                            "Trademark design phrase",
                            "conflict_grade",
                            "reasoning",
                        ]
                    )
                    # Create a table in the Word document
                    table_Name_Matchs = document.add_table(
                        df_Name_Matchs.shape[0] + 1, df_Name_Matchs.shape[1]
                    )
                    # Set a predefined table style (with borders)
                    table_Name_Matchs.style = (
                        "TableGrid"  # This is a built-in style that includes borders
                    )
                    # Add the column names to the table
                    for i, column_name in enumerate(df_Name_Matchs.columns):
                        table_Name_Matchs.cell(0, i).text = column_name
                    # Add the data to the table
                    for i, row in df_Name_Matchs.iterrows():
                        for j, value in enumerate(row):
                            cell = table_Name_Matchs.cell(i + 1, j)
                            cell.text = str(value)
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)

                if len(low_conflicts) > 0:
                    document.add_heading("Trademarks with 1 condition satisfied:", level=2)
                    # Create a pandas DataFrame from the JSON list
                    df_low = pd.DataFrame(low_conflicts)
                    df_low = df_low.drop(
                        columns=[
                            "Trademark name",
                            "Trademark class Number",
                            "Trademark registration number",
                            "Trademark serial number",
                            "Trademark design phrase",
                            "conflict_grade",
                            "reasoning",
                        ]
                    )
                    # Create a table in the Word document
                    table_low = document.add_table(df_low.shape[0] + 1, df_low.shape[1])
                    # Set a predefined table style (with borders)
                    table_low.style = (
                        "TableGrid"  # This is a built-in style that includes borders
                    )
                    # Add the column names to the table
                    for i, column_name in enumerate(df_low.columns):
                        table_low.cell(0, i).text = column_name
                    # Add the data to the table
                    for i, row in df_low.iterrows():
                        for j, value in enumerate(row):
                            cell = table_low.cell(i + 1, j)
                            cell.text = str(value)
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)

                def add_conflict_paragraph(document, conflict):
                    p = document.add_paragraph(
                        f"Trademark Name : {conflict.get('Trademark name', 'N/A')}"
                    )
                    p.paragraph_format.line_spacing = Pt(18)
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(
                        f"Trademark Status : {conflict.get('Trademark Status', 'N/A')}"
                    )
                    p.paragraph_format.line_spacing = Pt(18)
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(
                        f"Trademark Owner : {conflict.get('Trademark Owner', 'N/A')}"
                    )
                    p.paragraph_format.line_spacing = Pt(18)
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(
                        f"Trademark Class Number : {conflict.get('Trademark class Number', 'N/A')}"
                    )
                    p.paragraph_format.line_spacing = Pt(18)
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(
                        f"Trademark serial number : {conflict.get('Trademark serial number', 'N/A')}"
                    )
                    p.paragraph_format.line_spacing = Pt(18)
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(
                        f"Trademark registration number : {conflict.get('Trademark registration number', 'N/A')}"
                    )
                    p.paragraph_format.line_spacing = Pt(18)
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(
                        f"Trademark Design phrase : {conflict.get('Trademark design phrase', 'N/A')}"
                    )
                    p.paragraph_format.line_spacing = Pt(18)
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(" ")
                    p.paragraph_format.line_spacing = Pt(18)
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(f"{conflict.get('reasoning','N/A')}\n")
                    p.paragraph_format.line_spacing = Pt(18)
                    p = document.add_paragraph(" ")
                    p.paragraph_format.line_spacing = Pt(18)

                if len(high_conflicts) > 0:
                    document.add_heading(
                        "Explanation: Trademarks with 3 conditions satisfied:", level=2
                    )
                    p = document.add_paragraph(" ")
                    p.paragraph_format.line_spacing = Pt(18)
                    for conflict in high_conflicts:
                        add_conflict_paragraph(document, conflict)

                if len(moderate_conflicts) > 0:
                    document.add_heading(
                        "Explanation: Trademarks with 2 conditions satisfied:", level=2
                    )
                    p = document.add_paragraph(" ")
                    p.paragraph_format.line_spacing = Pt(18)
                    for conflict in moderate_conflicts:
                        add_conflict_paragraph(document, conflict)

                if len(Name_Matchs) > 0:
                    document.add_heading(
                        "Trademarks with Name Match's Conflicts Reasoning:", level=2
                    )
                    p = document.add_paragraph(" ")
                    p.paragraph_format.line_spacing = Pt(18)
                    for conflict in Name_Matchs:
                        add_conflict_paragraph(document, conflict)

                if len(low_conflicts) > 0:
                    document.add_heading(
                        "Explanation: Trademarks with 1 condition satisfied:", level=2
                    )
                    p = document.add_paragraph(" ")
                    p.paragraph_format.line_spacing = Pt(18)
                    for conflict in low_conflicts:
                        add_conflict_paragraph(document, conflict)



                def add_conflict_paragraph_to_array(conflict):  
                    result = []  
                    result.append(f"Trademark Name : {conflict.get('Trademark name', 'N/A')}")  
                    result.append(f"Trademark Status : {conflict.get('Trademark Status', 'N/A')}")  
                    result.append(f"Trademark Owner : {conflict.get('Trademark Owner', 'N/A')}")  
                    result.append(f"Trademark Class Number : {conflict.get('Trademark class Number', 'N/A')}")  
                    result.append(f"Trademark serial number : {conflict.get('Trademark serial number', 'N/A')}")  
                    result.append(f"Trademark registration number : {conflict.get('Trademark registration number', 'N/A')}")  
                    result.append(f"Trademark Design phrase : {conflict.get('Trademark design phrase', 'N/A')}")  
                    result.append(" ")  # Blank line for spacing  
                    result.append(f"{conflict.get('reasoning', 'N/A')}\n")  
                    result.append(" ")  # Blank line for spacing  
                    return result  
                
                conflicts_array = []  
                
                if len(high_conflicts) > 0:  
                    conflicts_array.append("Explanation: Trademarks with 3 conditions satisfied:")  
                    conflicts_array.append(" ")  # Blank line for spacing  
                    for conflict in high_conflicts:  
                        conflicts_array.extend(add_conflict_paragraph_to_array(conflict))  
                
                if len(moderate_conflicts) > 0:  
                    conflicts_array.append("Explanation: Trademarks with 2 conditions satisfied:")  
                    conflicts_array.append(" ")  # Blank line for spacing  
                    for conflict in moderate_conflicts:  
                        conflicts_array.extend(add_conflict_paragraph_to_array(conflict))  
                
                if len(Name_Matchs) > 0:  
                    conflicts_array.append("Trademarks with Name Match's Conflicts Reasoning:")  
                    conflicts_array.append(" ")  # Blank line for spacing  
                    for conflict in Name_Matchs:  
                        conflicts_array.extend(add_conflict_paragraph_to_array(conflict))  
                
                if len(low_conflicts) > 0:  
                    conflicts_array.append("Explanation: Trademarks with 1 condition satisfied:")  
                    conflicts_array.append(" ")  # Blank line for spacing  
                    for conflict in low_conflicts:  
                        conflicts_array.extend(add_conflict_paragraph_to_array(conflict))  
                    

                # for i in range(70,96):
                #     progress_bar.progress(i)


                progress_bar.progress(100)

                filename = proposed_name
                doc_stream = BytesIO()
                document.save(doc_stream)
                doc_stream.seek(0)
                download_table = f'<a href="data:application/octet-stream;base64,{base64.b64encode(doc_stream.read()).decode()}" download="{filename + " Trademark Conflict Report"}.docx">Download: {filename}</a>'
                st.sidebar.markdown(download_table, unsafe_allow_html=True)
                st.success(
                    f"{proposed_name} Document conflict report successfully completed!"
                )
                
                opinion_output = run_trademark_analysis(proposed_name, proposed_class, proposed_goods_services, conflicts_array)
                st.write("------------------------------------------------------------------------------------------------------------------------------")
                st.write(opinion_output)

                # Export to Word
                filename = export_trademark_opinion_to_word(opinion_output)
                
                # Download button
                with open(filename, "rb") as file:
                    st.sidebar.download_button(
                        label="Download Trademark Opinion",
                        data=file,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                end_time = time.time()
                elapsed_time = end_time - start_time
                elapsed_time = elapsed_time // 60
                st.write(f"Time taken: {elapsed_time} mins")

                st.write(
                    "______________________________________________________________________________________________________________________________"
                )

        progress_bar.progress(100)
        st.success("All documents processed successfully!")
